"""
文档服务
处理文档的上传、解析、向量化和存储
"""

import os
from pathlib import Path
from typing import List, Optional
from sqlmodel import Session, select

from dotenv import load_dotenv

from langchain_community.embeddings import DashScopeEmbeddings

from llama_index.core import VectorStoreIndex, SimpleDirectoryReader, Settings, StorageContext
from llama_index.embeddings.langchain import LangchainEmbedding
from llama_index.vector_stores.qdrant import QdrantVectorStore

from rag.models.models import Document, DocumentStatus
from rag.config import (
    SUPPORTED_EXTENSIONS,
    CHUNK_SIZE,
    CHUNK_OVERLAP,
    EMBEDDING_MODEL,
    get_document_storage_path,
)
from rag.service.qdrant_client import get_qdrant_client
from rag.utils.logger import get_logger

load_dotenv()

logger = get_logger(__name__)


class DocumentService:
    """文档服务"""

    def __init__(self):
        self._embed_model = None

    @property
    def client(self):
        """获取 Qdrant 客户端单例"""
        return get_qdrant_client()

    @property
    def embed_model(self):
        """获取嵌入模型"""
        if self._embed_model is None:
            logger.debug("初始化 Embedding 模型")
            embeddings = DashScopeEmbeddings(
                model=EMBEDDING_MODEL,
                dashscope_api_key=os.getenv("DASHSCOPE_API_KEY"),
            )
            self._embed_model = LangchainEmbedding(embeddings)
            Settings.embed_model = self._embed_model
            Settings.chunk_size = CHUNK_SIZE
            Settings.chunk_overlap = CHUNK_OVERLAP
            logger.info("Embedding 模型初始化完成")
        return self._embed_model

    def _get_file_type(self, filename: str) -> str:
        """获取文件类型"""
        ext = Path(filename).suffix.lower()
        return ext.lstrip(".") if ext else "unknown"

    def validate_file(self, filename: str) -> bool:
        """验证文件类型是否支持"""
        ext = Path(filename).suffix.lower()
        is_valid = ext in SUPPORTED_EXTENSIONS
        if not is_valid:
            logger.warning(f"不支持的文件类型：{ext}")
        return is_valid

    def save_upload_file(self, kb_id: int, filename: str, file_content: bytes) -> str:
        """保存上传的文件"""
        logger.debug(f"保存文件：kb_id={kb_id}, filename={filename}, size={len(file_content)} bytes")
        kb_dir = get_document_storage_path(kb_id)
        file_path = kb_dir / filename

        # 处理文件名冲突
        counter = 1
        while file_path.exists():
            stem = Path(filename).stem
            suffix = Path(filename).suffix
            file_path = kb_dir / f"{stem}_{counter}{suffix}"
            counter += 1
            logger.debug(f"文件名冲突，使用新名称：{file_path.name}")

        file_path.write_bytes(file_content)
        logger.info(f"文件保存成功：{file_path}")
        return str(file_path)

    def create_document_record(
        self,
        session: Session,
        kb_id: int,
        filename: str,
        file_path: str,
        file_size: int,
    ) -> Document:
        """创建文档记录"""
        logger.debug(f"创建文档记录：{filename}")
        doc = Document(
            kb_id=kb_id,
            filename=filename,
            file_path=file_path,
            file_size=file_size,
            file_type=self._get_file_type(filename),
            status=DocumentStatus.PENDING,
        )
        session.add(doc)
        session.commit()
        session.refresh(doc)
        logger.info(f"文档记录创建成功：{filename} (ID: {doc.id})")
        return doc

    def process_document(self, session: Session, doc: Document, collection_name: str) -> bool:
        """处理文档：解析并存储到向量数据库"""
        logger.info(f"开始处理文档：{doc.filename} (ID: {doc.id})")
        try:
            # 更新状态为处理中
            doc.status = DocumentStatus.PROCESSING
            session.add(doc)
            session.commit()
            logger.debug(f"文档状态更新为：PROCESSING")

            # 读取文档
            file_path = Path(doc.file_path)
            if not file_path.exists():
                error_msg = f"文件不存在：{file_path}"
                logger.error(error_msg)
                raise FileNotFoundError(error_msg)

            logger.debug(f"加载文档：{file_path}")
            # 根据文件类型选择加载方式
            documents = self._load_document(file_path)

            if not documents:
                error_msg = "文档内容为空"
                logger.error(error_msg)
                raise ValueError(error_msg)

            logger.info(f"文档加载成功，共 {len(documents)} 个文档块")

            # 创建向量存储并索引文档
            logger.debug(f"创建向量存储，集合：{collection_name}")
            vector_store = QdrantVectorStore(
                client=self.client,
                collection_name=collection_name,
            )
            storage_context = StorageContext.from_defaults(vector_store=vector_store)

            logger.info("开始向量化处理...")
            index = VectorStoreIndex.from_documents(
                documents,
                storage_context=storage_context,
                embed_model=self.embed_model,
            )
            logger.info(f"向量化处理完成，索引大小：{len(index.ref_doc_info)}")

            # 更新状态为完成
            doc.status = DocumentStatus.COMPLETED
            session.add(doc)
            session.commit()

            logger.info(f"文档处理成功：{doc.filename}")
            return True

        except Exception as e:
            # 更新状态为失败
            error_msg = str(e)[:500]
            logger.error(f"文档处理失败 [{doc.filename}]: {error_msg}", exc_info=True)
            doc.status = DocumentStatus.FAILED
            doc.error_message = error_msg
            session.add(doc)
            session.commit()
            return False

    def _load_document(self, file_path: Path):
        """根据文件类型加载文档"""
        ext = file_path.suffix.lower()
        logger.debug(f"加载文档，类型：{ext}")

        if ext in [".txt", ".md"]:
            # txt 和 md 使用 SimpleDirectoryReader
            logger.debug(f"使用 SimpleDirectoryReader 加载：{file_path}")
            return SimpleDirectoryReader(input_files=[str(file_path)]).load_data()

        elif ext == ".pdf":
            # PDF 文件
            logger.debug(f"加载 PDF 文件：{file_path}")
            return self._load_pdf(file_path)

        elif ext == ".docx":
            # Word 文档
            logger.debug(f"加载 Word 文档：{file_path}")
            return self._load_docx(file_path)

        else:
            error_msg = f"不支持的文件类型：{ext}"
            logger.error(error_msg)
            raise ValueError(error_msg)

    def _load_pdf(self, file_path: Path):
        """加载 PDF 文件"""
        from pypdf import PdfReader

        logger.debug(f"解析 PDF: {file_path}")
        reader = PdfReader(str(file_path))
        text = ""
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
            logger.debug(f"PDF 第 {i+1} 页解析完成")

        from llama_index.core import Document
        return [Document(text=text, metadata={"source": str(file_path)})]

    def _load_docx(self, file_path: Path):
        """加载 Word 文档"""
        from docx import Document as DocxDocument

        logger.debug(f"解析 Word 文档：{file_path}")
        doc = DocxDocument(str(file_path))
        text = "\n".join([para.text for para in doc.paragraphs if para.text])
        logger.debug(f"Word 文档解析完成，文本长度：{len(text)}")

        from llama_index.core import Document
        return [Document(text=text, metadata={"source": str(file_path)})]

    def list_documents(self, session: Session, kb_id: int) -> List[Document]:
        """获取知识库的文档列表"""
        logger.debug(f"获取文档列表：kb_id={kb_id}")
        statement = (
            select(Document)
            .where(Document.kb_id == kb_id)
            .order_by(Document.created_at.desc())
        )
        documents = session.exec(statement).all()
        logger.info(f"获取到 {len(documents)} 个文档")
        return documents

    def get_document(self, session: Session, doc_id: int) -> Optional[Document]:
        """获取单个文档"""
        logger.debug(f"获取文档：ID={doc_id}")
        return session.get(Document, doc_id)

    def delete_document(self, session: Session, doc_id: int) -> bool:
        """删除文档"""
        logger.info(f"删除文档：ID={doc_id}")
        try:
            doc = session.get(Document, doc_id)
            if not doc:
                logger.warning(f"文档不存在：ID={doc_id}")
                return False

            # 删除文件
            try:
                file_path = Path(doc.file_path)
                if file_path.exists():
                    logger.debug(f"删除文件：{file_path}")
                    file_path.unlink()
                    logger.info(f"文件删除成功：{file_path}")
            except Exception as e:
                logger.error(f"删除文件失败：{e}")

            # 删除数据库记录
            session.delete(doc)
            session.commit()

            logger.info(f"文档删除成功：{doc.filename} (ID: {doc_id})")
            return True

        except Exception as e:
            logger.error(f"删除文档失败：{e}", exc_info=True)
            return False


# 全局服务实例
document_service = DocumentService()